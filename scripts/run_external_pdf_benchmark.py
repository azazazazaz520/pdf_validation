from __future__ import annotations

import argparse
import json
import os
import sys
import time
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
os.environ.setdefault("PADDLE_PDX_CACHE_HOME", str(ROOT / "model_cache"))
sys.path.insert(0, str(ROOT))

from src.pdf_layout import extract_pdf_layout
from src.pdf_routing import analyze_pdf_text
from src.pdf_to_word_exporter import (
    export_results_to_docx,
    export_source_pages_to_docx,
    export_text_pages_to_docx,
)


SAMPLE_CATALOG: dict[str, dict[str, str]] = {
    "weknora_paper_test_report.pdf": {
        "category": "中文论文、图表、跨页表格",
        "source": "用户提供的本地样本",
    },
    "arxiv_attention_is_all_you_need.pdf": {
        "category": "英文双栏学术论文、公式、图表",
        "source": "https://arxiv.org/abs/1706.03762",
    },
    "who_phis_toolkit.pdf": {
        "category": "国际组织报告、图表、复杂版式",
        "source": "https://www.who.int/docs/default-source/documents/publications/public-health-information-services-toolkit.pdf",
    },
    "csrc_annual_report_guideline.pdf": {
        "category": "中文规范性文件、长文档、目录层级",
        "source": "https://www.csrc.gov.cn/ningxia/c105511/c7555161/content.shtml",
    },
    "novus_scanned_ocr_twin.pdf": {
        "category": "单页无文本层扫描件、OCR",
        "source": "https://examples.novusstreamsolutions.com/documents/pdf/scanned-ocr-twin-pdf",
    },
    "novus_simple_searchable.pdf": {
        "category": "单页可搜索 PDF、OCR 对照样本",
        "source": "https://examples.novusstreamsolutions.com/documents/pdf/simple-1-page-pdf",
    },
    "novus_bookmarked_toc_10_page.pdf": {
        "category": "多页、目录、书签、分页",
        "source": "https://examples.novusstreamsolutions.com/documents/pdf/bookmarked-toc-10-page",
    },
    "novus_landscape.pdf": {
        "category": "横向页面",
        "source": "https://examples.novusstreamsolutions.com/documents/pdf/landscape-pdf",
    },
    "novus_tables.pdf": {
        "category": "单页规则表格",
        "source": "https://examples.novusstreamsolutions.com/documents/pdf/tables-pdf",
    },
}


def _docx_metrics(path: Path) -> dict[str, Any]:
    document = Document(path)
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    full_text = "\n".join(paragraph_text + table_cells)
    headings = sum(
        paragraph.style.name.startswith("Heading")
        for paragraph in document.paragraphs
        if paragraph.style is not None
    )
    bullets = sum(
        paragraph.style.name in {"List Bullet", "List Bullet 2"}
        for paragraph in document.paragraphs
        if paragraph.style is not None
    )
    ordered = sum(
        paragraph.style.name in {"List Number", "List Number 2"}
        for paragraph in document.paragraphs
        if paragraph.style is not None
    )
    with zipfile.ZipFile(path) as archive:
        media_count = sum(name.startswith("word/media/") for name in archive.namelist())
        zip_ok = archive.testzip() is None
    return {
        "bytes": path.stat().st_size,
        "sections": len(document.sections),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "table_rows": sum(len(table.rows) for table in document.tables),
        "table_cells": len(table_cells),
        "headings": headings,
        "bullets": bullets,
        "ordered_lists": ordered,
        "text_chars": len("".join(full_text.split())),
        "media_count": media_count,
        "zip_ok": zip_ok,
    }


def _source_metrics(path: Path, min_page_chars: int) -> dict[str, Any]:
    reader = PdfReader(str(path))
    analysis = analyze_pdf_text(path, min_page_chars=min_page_chars)
    page_sizes = [
        [round(float(page.mediabox.width), 1), round(float(page.mediabox.height), 1)]
        for page in reader.pages
    ]
    layout_table_count = None
    layout_error = None
    try:
        layout_table_count = sum(len(page.tables) for page in extract_pdf_layout(path).pages)
    except Exception as error:
        layout_error = f"{type(error).__name__}: {error}"
    return {
        "bytes": path.stat().st_size,
        "pages": len(reader.pages),
        "page_sizes_pt": page_sizes,
        "text_chars": analysis.text_char_count,
        "usable_page_count": analysis.usable_page_count,
        "high_quality_page_count": analysis.high_quality_page_count,
        "full_page_image_page_count": analysis.full_page_image_page_count,
        "garbled_char_count": analysis.garbled_char_count,
        "image_page_count": sum(page.image_count > 0 for page in analysis.pages),
        "route": (
            "text"
            if analysis.has_complete_text_layer(
                min_page_chars=min_page_chars,
                min_high_quality_ratio=0.8,
            )
            else "page_image"
            if analysis.has_usable_text_layer(
                min_page_chars=min_page_chars,
                min_page_ratio=0.6,
            )
            else "ocr"
        ),
        "layout_table_count": layout_table_count,
        "layout_error": layout_error,
        "page_quality": [
            {
                "chars": page.text_char_count,
                "images": page.image_count,
                "full_page_images": page.full_page_image_count,
                "quality": page.quality_score,
            }
            for page in analysis.pages
        ],
    }


def _convert_text(path: Path, output_path: Path, min_page_chars: int) -> dict[str, Any]:
    analysis = analyze_pdf_text(path, min_page_chars=min_page_chars)
    started = time.perf_counter()
    events: list[dict[str, Any]] = []

    def callback(stage: str, details: dict[str, Any]) -> None:
        events.append({"stage": stage, **details})

    export_text_pages_to_docx(
        analysis.page_texts,
        output_path,
        title=path.stem,
        source_pdf=path,
        stage_callback=callback,
    )
    return {
        "status": "succeeded",
        "route": "text",
        "elapsed_seconds": round(time.perf_counter() - started, 3),
        "events": events,
    }


def _convert_page_image(path: Path, output_path: Path) -> dict[str, Any]:
    started = time.perf_counter()
    events: list[dict[str, Any]] = []

    def callback(stage: str, details: dict[str, Any]) -> None:
        events.append({"stage": stage, **details})

    export_source_pages_to_docx(
        output_path,
        title=path.stem,
        source_pdf=path,
        stage_callback=callback,
    )
    return {
        "status": "succeeded",
        "route": "page_image",
        "elapsed_seconds": round(time.perf_counter() - started, 3),
        "events": events,
    }


def _convert_ocr(path: Path, output_path: Path, pipeline: Any) -> dict[str, Any]:
    started = time.perf_counter()
    results = list(pipeline.predict_iter(str(path)))
    ocr_pages: list[dict[str, Any]] = []
    for result in results:
        overall = result.get("overall_ocr_res", {}) if isinstance(result, dict) else {}
        texts = overall.get("rec_texts", []) if isinstance(overall, dict) else []
        scores = overall.get("rec_scores", []) if isinstance(overall, dict) else []
        ocr_pages.append(
            {
                "line_count": len(texts),
                "text_chars": len("".join(str(text) for text in texts)),
                "mean_confidence": round(sum(float(score) for score in scores) / len(scores), 4)
                if scores
                else None,
            }
        )
    export_results_to_docx(
        results,
        output_path,
        title=path.stem,
        source_pdf=path,
        mode="hybrid",
    )
    return {
        "status": "succeeded",
        "route": "ocr",
        "elapsed_seconds": round(time.perf_counter() - started, 3),
        "ocr_result_count": len(results),
        "ocr_pages": ocr_pages,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="批量验证公开 PDF 转 Word 效果")
    parser.add_argument(
        "--source-dir",
        type=Path,
        default=ROOT / "artifacts" / "test_samples" / "sources",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=ROOT / "artifacts" / "test_samples" / "results",
    )
    parser.add_argument(
        "--rendered-dir",
        type=Path,
        default=ROOT / "artifacts" / "test_samples" / "rendered",
    )
    parser.add_argument("--with-ocr", action="store_true")
    parser.add_argument("--min-page-chars", type=int, default=20)
    args = parser.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    results: list[dict[str, Any]] = []
    pipeline = None
    for source_path in sorted(args.source_dir.glob("*.pdf")):
        started = time.perf_counter()
        record: dict[str, Any] = {
            "name": source_path.name,
            "category": SAMPLE_CATALOG.get(source_path.name, {}).get("category", "未分类"),
            "source": SAMPLE_CATALOG.get(source_path.name, {}).get("source", "本地样本"),
        }
        try:
            record["source_metrics"] = _source_metrics(source_path, args.min_page_chars)
            route = record["source_metrics"]["route"]
            output_path = args.output_dir / f"{source_path.stem}.docx"
            if route == "text":
                conversion = _convert_text(source_path, output_path, args.min_page_chars)
            elif route == "page_image":
                conversion = _convert_page_image(source_path, output_path)
            elif args.with_ocr:
                if pipeline is None:
                    from src.run_validation import build_pipeline

                    pipeline = build_pipeline("structure-lite")
                conversion = _convert_ocr(source_path, output_path, pipeline)
            else:
                conversion = {
                    "status": "skipped",
                    "route": "ocr",
                    "reason": "未启用 --with-ocr，避免在未确认模型资源时启动 OCR",
                }
            record["conversion"] = conversion
            if output_path.is_file():
                record["output_metrics"] = _docx_metrics(output_path)
            rendered_path = args.rendered_dir / f"{source_path.stem}.pdf"
            if rendered_path.is_file():
                record["rendered_pdf_pages"] = len(PdfReader(str(rendered_path)).pages)
                record["rendered_page_delta"] = (
                    record["rendered_pdf_pages"] - record["source_metrics"]["pages"]
                )
        except Exception as error:
            record["conversion"] = {
                "status": "failed",
                "error": f"{type(error).__name__}: {error}",
            }
        record["elapsed_seconds_total"] = round(time.perf_counter() - started, 3)
        results.append(record)
        print(json.dumps(record, ensure_ascii=False), flush=True)

    report_path = args.output_dir / "benchmark.json"
    report_path.write_text(
        json.dumps(
            {
                "status": "completed",
                "with_ocr": args.with_ocr,
                "python": sys.version,
                "results": results,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )
    print(json.dumps({"report": str(report_path), "count": len(results)}, ensure_ascii=False))
    return 0 if all(item.get("conversion", {}).get("status") != "failed" for item in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
