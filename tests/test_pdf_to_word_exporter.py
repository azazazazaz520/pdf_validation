from __future__ import annotations

from io import BytesIO
import zipfile
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.oxml.ns import qn
from PIL import Image
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen.canvas import Canvas

from src.pdf_layout import extract_pdf_layout
from src.pdf_to_word_exporter import (
    _iter_page_images,
    _add_content_lines,
    _add_text_pages_with_sizes,
    _group_text_lines,
    _set_document_styles,
    export_results_to_docx,
    export_source_pages_to_docx,
    export_text_pages_to_docx,
)


class HybridExportTest(unittest.TestCase):
    def test_layout_export_preserves_headings_bullets_and_vector_table(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            pdf_path = root / "layout.pdf"
            docx_path = root / "layout.docx"

            canvas = Canvas(str(pdf_path), pagesize=(360, 300))
            canvas.setFont("Helvetica", 16)
            canvas.drawString(36, 270, "Layout title")
            canvas.setFont("Helvetica", 12)
            canvas.drawString(36, 250, "1. Scope")
            canvas.drawString(60, 230, "First bullet.")
            canvas.drawString(60, 212, "Second bullet.")
            for y in (80, 110, 140, 170):
                canvas.line(36, y, 324, y)
            for x in (36, 180, 324):
                canvas.line(x, 80, x, 170)
            canvas.setFont("Helvetica", 10)
            canvas.drawString(45, 150, "Name")
            canvas.drawString(190, 150, "Value")
            canvas.drawString(45, 120, "Alpha")
            canvas.drawString(190, 120, "A")
            canvas.drawString(45, 90, "Beta")
            canvas.drawString(190, 90, "B")
            canvas.save()

            layout = extract_pdf_layout(pdf_path)
            self.assertEqual(layout.table_count, 1)
            self.assertEqual(layout.pages[0].tables[0].row_count, 3)
            self.assertEqual(layout.pages[0].tables[0].column_count, 2)
            self.assertEqual(layout.pages[0].tables[0].rows[0], ("Name", "Value"))

            export_text_pages_to_docx(
                ["ignored"],
                docx_path,
                title="layout-test",
                source_pdf=pdf_path,
            )

            document = Document(str(docx_path))
            headings = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style.name.startswith("Heading")
            ]
            self.assertEqual(headings[:2], ["Layout title", "1. Scope"])
            self.assertEqual(
                [
                    paragraph.text
                    for paragraph in document.paragraphs
                    if paragraph.style.name == "List Bullet"
                ],
                ["First bullet.", "Second bullet."],
            )
            self.assertEqual(len(document.tables), 1)
            self.assertEqual(document.tables[0].cell(2, 1).text, "B")

    def test_layout_export_preserves_merged_cells_and_row_heights(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            pdf_path = root / "merged-table.pdf"
            docx_path = root / "merged-table.docx"

            canvas = Canvas(str(pdf_path), pagesize=(360, 300))
            horizontal_lines = (
                (36, 220, 284),
                (36, 180, 284),
                (36, 140, 160),
                (36, 100, 284),
                (36, 60, 284),
            )
            for x0, y, x1 in horizontal_lines:
                canvas.line(x0, y, x1, y)
            canvas.line(36, 60, 36, 220)
            canvas.line(160, 60, 160, 180)
            canvas.line(284, 60, 284, 220)
            canvas.setFont("Helvetica", 10)
            canvas.drawString(45, 200, "Merged header")
            canvas.drawString(45, 160, "Group")
            canvas.drawString(170, 160, "Shared value")
            canvas.drawString(45, 120, "A")
            canvas.drawString(45, 80, "B")
            canvas.drawString(170, 80, "Final value")
            canvas.save()

            layout = extract_pdf_layout(pdf_path)
            pdf_table = layout.pages[0].tables[0]
            cells = {
                (cell.row_index, cell.column_index): cell
                for cell in pdf_table.cells
            }
            self.assertEqual(pdf_table.row_count, 4)
            self.assertEqual(pdf_table.column_count, 2)
            self.assertEqual(pdf_table.row_heights, (40.0, 40.0, 40.0, 40.0))
            self.assertEqual(cells[(0, 0)].column_span, 2)
            self.assertEqual(cells[(1, 1)].row_span, 2)
            self.assertEqual(cells[(0, 0)].text, "Merged header")
            self.assertEqual(cells[(1, 1)].text, "Shared value")

            export_text_pages_to_docx(
                ["ignored"],
                docx_path,
                title="merged-table-test",
                source_pdf=pdf_path,
            )

            table = Document(str(docx_path)).tables[0]
            self.assertFalse(table.autofit)
            self.assertGreater(table.rows[0].height.pt, 0)
            self.assertEqual(
                table._tbl.tblPr.find(qn("w:tblLayout")).get(qn("w:type")),
                "fixed",
            )
            self.assertIsNotNone(table.rows[0]._tr.trPr.find(qn("w:tblHeader")))
            self.assertIsNotNone(table.rows[0]._tr.trPr.find(qn("w:cantSplit")))
            self.assertIs(table.cell(0, 0)._tc, table.cell(0, 1)._tc)
            self.assertIs(table.cell(1, 1)._tc, table.cell(2, 1)._tc)
            self.assertEqual(table.cell(0, 0).text, "Merged header")
            self.assertEqual(table.cell(1, 1).text, "Shared value")

    def test_layout_export_repeats_header_for_cross_page_table(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            pdf_path = root / "continued-table.pdf"
            docx_path = root / "continued-table.docx"

            canvas = Canvas(str(pdf_path), pagesize=(360, 300))
            for y in (280, 240, 200, 160, 120, 80, 20):
                canvas.line(36, y, 284, y)
            for x in (36, 160, 284):
                canvas.line(x, 20, x, 280)
            canvas.setFont("Helvetica", 10)
            for x, text in ((45, "Name"), (170, "Value")):
                canvas.drawString(x, 260, text)
            for y, name, value in (
                (220, "Alpha", "A"),
                (180, "Beta", "B"),
                (140, "Delta", "D"),
                (100, "Epsilon", "E"),
                (60, "Zeta", "Z"),
            ):
                canvas.drawString(45, y, name)
                canvas.drawString(170, y, value)
            canvas.showPage()

            for y in (280, 240, 200):
                canvas.line(36, y, 284, y)
            for x in (36, 160, 284):
                canvas.line(x, 200, x, 280)
            for y, name, value in ((260, "Gamma", "C"), (220, "Eta", "G")):
                canvas.drawString(45, y, name)
                canvas.drawString(170, y, value)
            canvas.save()

            layout = extract_pdf_layout(pdf_path)
            self.assertEqual(layout.table_count, 2)
            continuation = layout.pages[1].tables[0]
            self.assertTrue(continuation.continued_from_previous_page)
            self.assertEqual(continuation.continuation_header_rows, (("Name", "Value"),))

            export_text_pages_to_docx(
                ["ignored", "ignored"],
                docx_path,
                title="continued-table-test",
                source_pdf=pdf_path,
            )

            document = Document(str(docx_path))
            self.assertEqual(len(document.tables), 2)
            continued_table = document.tables[1]
            self.assertEqual(continued_table.cell(0, 0).text, "Name")
            self.assertEqual(continued_table.cell(0, 1).text, "Value")
            self.assertEqual(continued_table.cell(1, 0).text, "Gamma")
            self.assertEqual(continued_table.cell(2, 1).text, "G")

    def test_text_grouping_joins_wrapped_lines_without_extra_spaces(self) -> None:
        blocks = _group_text_lines(
            "标题\n第一行内容\n第二行内容。\n1. 第一项\n续行。",
            is_document_start=True,
        )

        self.assertEqual(
            blocks,
            [
                ("title", "标题"),
                ("body", "第一行内容第二行内容。"),
                ("ordered", "第一项续行。"),
            ],
        )

    def test_list_item_colon_keeps_following_explanation_in_same_item(self) -> None:
        blocks = _group_text_lines(
            "四、总结\n1. 第一项：\n这是第一项的说明。\n2. 第二项：\n这是第二项的说明。"
        )

        self.assertEqual(
            blocks,
            [
                ("heading1", "四、总结"),
                ("ordered", "第一项：这是第一项的说明。"),
                ("ordered", "第二项：这是第二项的说明。"),
            ],
        )

    def test_common_parenthesized_numbers_are_ordered_items(self) -> None:
        blocks = _group_text_lines("（1）第一项。\n(2) 第二项。")

        self.assertEqual(
            blocks,
            [("ordered", "第一项。"), ("ordered", "第二项。")],
        )

    def test_common_math_operator_line_is_a_formula(self) -> None:
        blocks = _group_text_lines("目标函数：\n∑ᵢ xᵢ = 1")

        self.assertEqual(
            blocks,
            [("body", "目标函数："), ("formula", "∑ᵢ xᵢ = 1")],
        )

    def test_ocr_content_recovers_heading_and_restarts_list(self) -> None:
        document = Document()
        _set_document_styles(document)

        _add_content_lines(
            document,
            "1. 第一项：\n这是第一项的说明。\n参考文献\n1. 文献",
            "text",
        )

        paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
        self.assertEqual(paragraphs[0].style.name, "List Number")
        self.assertEqual(paragraphs[0].text, "第一项：这是第一项的说明。")
        self.assertEqual(paragraphs[1].style.name, "Heading 1")
        self.assertEqual(paragraphs[2].style.name, "List Number")

    def test_ordered_lists_restart_after_a_heading(self) -> None:
        document = Document()
        _set_document_styles(document)
        _add_text_pages_with_sizes(
            document,
            ["1. 第一项\n2. 第二项", "3. 第三项", "参考文献\n1. 文献"],
            None,
        )

        list_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name == "List Number"
        ]
        number_ids = [
            paragraph._p.pPr.numPr.numId.val for paragraph in list_paragraphs
        ]
        self.assertEqual(number_ids[0], number_ids[1])
        self.assertEqual(number_ids[1], number_ids[2])
        self.assertNotEqual(number_ids[2], number_ids[3])

    def test_bullet_items_use_bullet_numbering_definition(self) -> None:
        document = Document()
        _set_document_styles(document)
        _add_text_pages_with_sizes(document, ["• 第一项\n• 第二项"], None)

        bullet_paragraph = next(
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style.name == "List Bullet"
        )
        num_id = bullet_paragraph._p.pPr.numPr.numId.val
        numbering = document.part.numbering_part.element
        num = next(
            element
            for element in numbering.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")
            if element.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId")
            == str(num_id)
        )
        abstract_id = num.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId"
        ).get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
        abstract = next(
            element
            for element in numbering.findall(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNum"
            )
            if element.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}abstractNumId"
            )
            == abstract_id
        )
        num_format = abstract.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numFmt"
        )
        self.assertEqual(
            num_format.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"),
            "bullet",
        )

    def test_decimal_section_heading_is_not_numbered_as_a_list(self) -> None:
        document = Document()
        _add_content_lines(document, "2.1 原始线性规划模型\n1. 普通编号条目", "text")

        paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text]
        self.assertEqual(paragraphs[0].style.name, "Heading 2")
        self.assertEqual(paragraphs[1].style.name, "List Number")

    def test_text_export_uses_source_page_size(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            pdf_path = root / "a4-input.pdf"
            docx_path = root / "output.docx"

            canvas = Canvas(str(pdf_path), pagesize=(595.28, 841.89))
            canvas.drawString(72, 760, "Editable page text")
            canvas.save()

            export_text_pages_to_docx(
                ["Editable page text"],
                docx_path,
                title="text-size-test",
                source_pdf=pdf_path,
            )

            document = Document(str(docx_path))
            section = document.sections[0]
            self.assertAlmostEqual(section.page_width.inches, 595.28 / 72, places=2)
            self.assertAlmostEqual(section.page_height.inches, 841.89 / 72, places=2)

    def test_hybrid_export_keeps_page_image_and_ocr_text(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            pdf_path = root / "input.pdf"
            docx_path = root / "output.docx"

            canvas = Canvas(str(pdf_path), pagesize=(360, 240))
            canvas.drawString(36, 200, "Rendered PDF page")
            canvas.rect(36, 36, 288, 120)
            canvas.save()

            results = [
                {
                    "parsing_res_list": [
                        {
                            "block_label": "title",
                            "block_content": "示例 OCR 文本",
                            "block_bbox": [36, 36, 240, 80],
                            "block_order": 0,
                        }
                    ]
                }
            ]

            export_results_to_docx(
                results,
                docx_path,
                title="hybrid-test",
                source_pdf=pdf_path,
                mode="hybrid",
            )

            document = Document(str(docx_path))
            paragraphs = "\n".join(p.text for p in document.paragraphs)
            self.assertIn("示例 OCR 文本", paragraphs)
            self.assertNotIn("OCR 可编辑文本", paragraphs)
            self.assertNotIn("以下文字由 OCR 识别", paragraphs)
            self.assertGreaterEqual(len(document.sections), 2)
            with zipfile.ZipFile(docx_path) as archive:
                media = [name for name in archive.namelist() if name.startswith("word/media/")]
            self.assertEqual(len(media), 1)

    def test_page_render_keeps_vector_overlay_on_full_page_image(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            pdf_path = root / "image-with-overlay.pdf"
            background_path = root / "background.png"
            Image.new("RGB", (720, 480), color="white").save(background_path)
            canvas = Canvas(str(pdf_path), pagesize=(360, 240))
            canvas.drawImage(
                ImageReader(str(background_path)),
                0,
                0,
                width=360,
                height=240,
            )
            canvas.setFillColorRGB(0, 0, 0)
            canvas.drawString(36, 200, "Vector overlay text")
            canvas.save()

            page = next(
                _iter_page_images(
                    pdf_path,
                    max_pixels=4 * 1024 * 1024,
                    jpeg_quality=88,
                )
            )
            with Image.open(BytesIO(page[1])) as rendered:
                rgb = rendered.convert("RGB")
                crop = rgb.crop((50, 40, 400, 160))
                pixels = crop.load()
                dark_pixels = sum(
                    max(pixels[x, y]) < 100
                    for x in range(crop.width)
                    for y in range(crop.height)
                )
            self.assertGreater(dark_pixels, 10)

    def test_source_page_export_does_not_append_white_text_pages(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            pdf_path = root / "input.pdf"
            docx_path = root / "output.docx"

            canvas = Canvas(str(pdf_path), pagesize=(360, 240))
            canvas.drawString(36, 200, "Rendered PDF page")
            canvas.rect(36, 36, 288, 120)
            canvas.save()

            export_source_pages_to_docx(
                docx_path,
                title="source-page-test",
                source_pdf=pdf_path,
            )

            document = Document(str(docx_path))
            paragraphs = "\n".join(p.text for p in document.paragraphs)
            self.assertNotIn("Selectable text layer", paragraphs)
            self.assertEqual(len(document.sections), 1)
            with zipfile.ZipFile(docx_path) as archive:
                media = [name for name in archive.namelist() if name.startswith("word/media/")]
            self.assertEqual(len(media), 1)

    def test_ocr_text_export_uses_source_page_size(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            root = Path(temporary_directory)
            pdf_path = root / "a4-input.pdf"
            docx_path = root / "output.docx"

            canvas = Canvas(str(pdf_path), pagesize=(595.28, 841.89))
            canvas.drawString(72, 760, "OCR page text")
            canvas.save()

            export_results_to_docx(
                [{"parsing_res_list": [{"block_content": "OCR page text"}]}],
                docx_path,
                title="ocr-size-test",
                source_pdf=pdf_path,
                mode="text",
            )

            section = Document(str(docx_path)).sections[0]
            self.assertAlmostEqual(section.page_width.inches, 595.28 / 72, places=2)
            self.assertAlmostEqual(section.page_height.inches, 841.89 / 72, places=2)


if __name__ == "__main__":
    unittest.main()
