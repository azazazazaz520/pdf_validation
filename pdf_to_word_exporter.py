from __future__ import annotations

import math
import re
from io import BytesIO
from html.parser import HTMLParser
from pathlib import Path
from collections.abc import Callable, Iterable, Mapping
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from PIL import Image
from pypdf import PdfReader


_ORDERED_ITEM = re.compile(
    r"^\s*(?:(?:\d+[.．、)](?!\d))|(?:[（(]\s*\d+\s*[)）])|(?:[①-⑳]))\s*(.+)$"
)
_BULLET_ITEM = re.compile(r"^\s*[•·●▪◦]\s*(.+)$")
_PLUGIN_ITEM = re.compile(r"^\s*(plugin_[a-zA-Z0-9_]+)\s*[:：]\s*(.+)$")
_SECTION_HEADING = re.compile(
    r"^\s*(?:[一二三四五六七八九十百]+、|摘要(?:\s|$)|参考文献(?:\s|（|\(|$))"
)
_SUBSECTION_HEADING = re.compile(r"^\s*\d+\.\d+(?:\s|$)")
_FORMULA_PREFIX = re.compile(r"^\s*(?:max|min|s\.\s*t\.)(?:\s|$)", re.IGNORECASE)
_FORMULA_SYMBOL_PREFIX = re.compile(r"^\s*[∑∏∫√∞∂∇∀∃]+")
_LIST_ROLES = frozenset({"ordered", "bullet", "plugin"})
_HTML_TAG = re.compile(r"<[^>]+>")
EXPORT_MODES = frozenset({"hybrid", "text"})
DEFAULT_PAGE_IMAGE_MAX_PIXELS = 4 * 1024 * 1024
DEFAULT_PAGE_IMAGE_JPEG_QUALITY = 88
MAX_PAGE_DIMENSION_INCHES = 11.0

StageCallback = Callable[[str, dict[str, Any]], None]


class _TableParser(HTMLParser):
    """读取 PaddleOCR 表格 HTML 中的行和单元格文本。"""

    def __init__(self) -> None:
        super().__init__()
        self.rows: list[list[str]] = []
        self._row: list[str] | None = None
        self._cell: list[str] | None = None

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag == "tr":
            self._row = []
        elif tag in ("td", "th"):
            self._cell = []

    def handle_endtag(self, tag: str) -> None:
        if tag in ("td", "th") and self._row is not None and self._cell is not None:
            self._row.append("".join(self._cell).strip())
            self._cell = None
        elif tag == "tr" and self._row is not None:
            if self._row:
                self.rows.append(self._row)
            self._row = None

    def handle_data(self, data: str) -> None:
        if self._cell is not None:
            self._cell.append(data)


def _plain_text(value: str) -> str:
    value = _HTML_TAG.sub("", value)
    return re.sub(r"[ \t]+", " ", value).strip()


def _table_rows(value: str) -> list[list[str]]:
    parser = _TableParser()
    parser.feed(value)
    return parser.rows


def _field(value: Any, name: str, default: Any = None) -> Any:
    if isinstance(value, Mapping):
        return value.get(name, default)
    attribute = getattr(value, name, None)
    if attribute is not None:
        return attribute
    aliases = {
        "block_label": "label",
        "block_content": "content",
        "block_bbox": "bbox",
        "block_order": "order_index",
    }
    return getattr(value, aliases.get(name, name), default)


def _ordered_blocks(result: Any) -> Iterable[Any]:
    blocks = list(_field(result, "parsing_res_list", []) or [])
    indexed = list(enumerate(blocks))

    def sort_key(item: tuple[int, Any]) -> tuple[int, int]:
        index, block = item
        order = _field(block, "block_order")
        if isinstance(order, (int, float)):
            return (0, int(order))
        bbox = _field(block, "block_bbox") or [0, index]
        return (1, int(bbox[1]) if len(bbox) > 1 else index)

    for _, block in sorted(indexed, key=sort_key):
        yield block


def _set_document_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    for style_name in ("Normal", "List Paragraph", "List Bullet", "List Number"):
        style = document.styles[style_name]
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.bold = True
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1


def _notify_stage(
    callback: StageCallback | None, stage: str, **details: Any
) -> None:
    if callback is not None:
        callback(stage, details)


def _set_section_page(
    section: Any,
    width_inches: float,
    height_inches: float,
    margin_inches: float,
) -> None:
    from docx.shared import Inches

    section.page_width = Inches(width_inches)
    section.page_height = Inches(height_inches)
    section.top_margin = Inches(margin_inches)
    section.bottom_margin = Inches(margin_inches)
    section.left_margin = Inches(margin_inches)
    section.right_margin = Inches(margin_inches)


def _scaled_page_size(
    width_points: float, height_points: float
) -> tuple[float, float]:
    width_inches = max(width_points / 72.0, 0.01)
    height_inches = max(height_points / 72.0, 0.01)
    scale = min(1.0, MAX_PAGE_DIMENSION_INCHES / max(width_inches, height_inches))
    return width_inches * scale, height_inches * scale


def _encode_page_image(
    image: Image.Image,
    *,
    max_pixels: int,
    jpeg_quality: int,
) -> bytes:
    image = image.convert("RGB")
    pixel_count = image.width * image.height
    if pixel_count > max_pixels:
        scale = math.sqrt(max_pixels / pixel_count)
        image = image.resize(
            (max(1, round(image.width * scale)), max(1, round(image.height * scale))),
            Image.Resampling.LANCZOS,
        )
    output = BytesIO()
    image.save(
        output,
        format="JPEG",
        quality=max(1, min(100, jpeg_quality)),
        optimize=True,
    )
    return output.getvalue()


def _render_page_image(
    pdf_document: Any,
    page_index: int,
    *,
    max_pixels: int,
    jpeg_quality: int,
) -> bytes:
    import pypdfium2 as pdfium

    pdf_page = pdf_document[page_index]
    try:
        width, height = pdf_page.get_size()
        scale = min(2.0, math.sqrt(max_pixels / max(width * height, 1.0)))
        bitmap = pdf_page.render(scale=max(scale, 0.1))
        try:
            image = bitmap.to_pil()
            return _encode_page_image(
                image,
                max_pixels=max_pixels,
                jpeg_quality=jpeg_quality,
            )
        finally:
            close_bitmap = getattr(bitmap, "close", None)
            if close_bitmap is not None:
                close_bitmap()
    finally:
        pdf_page.close()


def _iter_page_images(
    source_pdf: Path,
    *,
    max_pixels: int,
    jpeg_quality: int,
) -> Iterable[tuple[int, bytes, float, float]]:
    import pypdfium2 as pdfium

    reader = PdfReader(str(source_pdf))
    pdf_document = pdfium.PdfDocument(str(source_pdf))
    try:
        for page_index, pdf_page in enumerate(reader.pages):
            page_width = float(pdf_page.mediabox.width)
            page_height = float(pdf_page.mediabox.height)
            image_bytes = _render_page_image(
                pdf_document,
                page_index,
                max_pixels=max_pixels,
                jpeg_quality=jpeg_quality,
            )
            yield page_index, image_bytes, page_width, page_height
    finally:
        pdf_document.close()


def _add_page_image(
    document: Document,
    section: Any,
    image_bytes: bytes,
) -> None:
    paragraph = (
        document.paragraphs[0]
        if len(document.paragraphs) == 1
        else document.add_paragraph()
    )
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.add_run().add_picture(BytesIO(image_bytes), width=section.page_width)


def _add_source_pages(
    document: Document,
    source_pdf: Path,
    *,
    max_pixels: int,
    jpeg_quality: int,
    stage_callback: StageCallback | None,
) -> int:
    page_count = 0
    for page_index, image_bytes, page_width, page_height in _iter_page_images(
        source_pdf,
        max_pixels=max_pixels,
        jpeg_quality=jpeg_quality,
    ):
        if page_index > 0:
            section = document.add_section(WD_SECTION.NEW_PAGE)
        else:
            section = document.sections[0]
        width_inches, height_inches = _scaled_page_size(page_width, page_height)
        _set_section_page(section, width_inches, height_inches, 0)
        _add_page_image(
            document,
            section,
            image_bytes,
        )
        page_count += 1
        _notify_stage(
            stage_callback,
            "page_image_completed",
            page=page_index + 1,
            page_count=page_count,
        )
    return page_count


def _add_table(document: Document, html: str) -> None:
    rows = _table_rows(html)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            if index < len(cells):
                cells[index].text = value


def _add_content_lines(
    document: Document,
    content: str,
    block_label: str,
    *,
    list_state: dict[str, int | None | str] | None = None,
) -> None:
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    if not lines:
        return

    if block_label == "abstract" and len(lines[0]) <= 24 and not re.search(r"[：:。；;]", lines[0]):
        document.add_heading(_plain_text(lines.pop(0)), level=1)
        if list_state is not None:
            list_state["kind"] = None
            list_state["num_id"] = None
    if not lines:
        return

    state = list_state or {"kind": None, "num_id": None}
    for role, text in _group_text_lines("\n".join(lines)):
        _add_editable_text_block(document, role, _plain_text(text), state)


def _is_formula_line(line: str) -> bool:
    if _FORMULA_PREFIX.match(line) or _FORMULA_SYMBOL_PREFIX.match(line):
        return True
    if not re.match(r"^\s*(?:[A-Za-zα-ωΑ-Ω]|\d)", line):
        return False
    return bool(
        re.search(r"[=≤≥≈]", line)
        and re.search(r"[A-Za-zα-ωΑ-Ω][₀-₉]?", line)
    )


def _text_line_role(line: str, *, is_first_line: bool) -> str:
    if _SECTION_HEADING.match(line):
        return "heading1"
    if _SUBSECTION_HEADING.match(line):
        return "heading2"
    if _ORDERED_ITEM.match(line):
        return "ordered"
    if _BULLET_ITEM.match(line):
        return "bullet"
    if _PLUGIN_ITEM.match(line):
        return "plugin"
    if _is_formula_line(line):
        return "formula"
    if is_first_line:
        return "title"
    return "body"


def _is_sentence_terminal(line: str) -> bool:
    return bool(re.search(r"[。！？!?；;：:]$", line))


def _needs_word_space(previous: str, current: str) -> bool:
    if not previous or not current:
        return False
    if previous[-1] in "+-=*/≤≥≈":
        return True
    return previous[-1].isascii() and current[0].isascii() and (
        previous[-1].isalnum() and current[0].isalnum()
    )


def _join_wrapped_lines(lines: list[str]) -> str:
    if not lines:
        return ""
    value = lines[0]
    for line in lines[1:]:
        if _needs_word_space(value, line):
            value += " "
        value += line
    return value


def _group_text_lines(
    content: str,
    *,
    is_document_start: bool = False,
) -> list[tuple[str, str]]:
    """将 PDF 物理换行合并为逻辑段落，并识别标题、列表和公式。"""
    lines = [line.strip() for line in content.splitlines() if line.strip()]
    blocks: list[tuple[str, str]] = []
    current_role: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        nonlocal current_role, current_lines
        if current_role is not None and current_lines:
            text = _join_wrapped_lines(current_lines)
            if current_role == "ordered":
                match = _ORDERED_ITEM.match(text)
                text = match.group(1) if match else text
            elif current_role == "bullet":
                match = _BULLET_ITEM.match(text)
                text = match.group(1) if match else text
            elif current_role == "plugin":
                match = _PLUGIN_ITEM.match(text)
                if match:
                    text = f"{match.group(1)}：{match.group(2)}"
            blocks.append((current_role, text))
        current_role = None
        current_lines = []

    for index, line in enumerate(lines):
        role = _text_line_role(
            line,
            is_first_line=is_document_start and index == 0,
        )
        if role in {"title", "heading1", "heading2", "formula"}:
            flush()
            blocks.append((role, line))
            continue

        if role in _LIST_ROLES:
            flush()
            current_role = role
            current_lines = [line]
            if _is_sentence_terminal(line) and not line.endswith(("：", ":")):
                flush()
            continue

        if current_role in _LIST_ROLES:
            current_lines.append(line)
            if _is_sentence_terminal(line) and not line.endswith(("：", ":")):
                flush()
            continue

        if current_role != "body":
            current_role = "body"
            current_lines = []
        current_lines.append(line)
        if _is_sentence_terminal(line):
            flush()

    flush()
    return blocks


def _create_numbering_instance(document: Document, *, ordered: bool = True) -> int:
    numbering = document.part.numbering_part.element
    num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
        if element.get(qn("w:numId"))
    ]
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
        if element.get(qn("w:abstractNumId"))
    ]
    num_id = max(num_ids, default=0) + 1
    abstract_id = max(abstract_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id:08X}")
    abstract.append(nsid)
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)
    template = OxmlElement("w:tmpl")
    template.set(qn("w:val"), f"{(abstract_id + 1):08X}")
    abstract.append(template)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_format)
    paragraph_style = OxmlElement("w:pStyle")
    paragraph_style.set(qn("w:val"), "ListNumber" if ordered else "ListBullet")
    level.append(paragraph_style)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "\uf0b7")
    level.append(level_text)
    alignment = OxmlElement("w:lvlJc")
    alignment.set(qn("w:val"), "left")
    level.append(alignment)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "720")
    indentation.set(qn("w:hanging"), "360")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    if not ordered:
        run_properties = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        fonts.set(qn("w:hint"), "default")
        run_properties.append(fonts)
        level.append(run_properties)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    numbering_instance = OxmlElement("w:num")
    numbering_instance.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    numbering_instance.append(abstract_reference)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    numbering_instance.append(level_override)
    numbering.append(numbering_instance)
    return num_id


def _set_numbering_instance(paragraph: Any, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering_properties = paragraph_properties.find(qn("w:numPr"))
    if numbering_properties is None:
        numbering_properties = OxmlElement("w:numPr")
        paragraph_properties.append(numbering_properties)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    numbering_properties.append(level)
    numbering_id = OxmlElement("w:numId")
    numbering_id.set(qn("w:val"), str(num_id))
    numbering_properties.append(numbering_id)


def _add_editable_text_block(
    document: Document,
    role: str,
    text: str,
    list_state: dict[str, int | None | str],
) -> None:
    if role == "ordered":
        if list_state.get("kind") != "ordered" or list_state.get("num_id") is None:
            list_state["num_id"] = _create_numbering_instance(document)
        paragraph = document.add_paragraph(style="List Number")
        _set_numbering_instance(paragraph, int(list_state["num_id"]))
        paragraph.add_run(text)
        list_state["kind"] = "ordered"
        return
    if role in {"bullet", "plugin"}:
        if list_state.get("kind") != "bullet" or list_state.get("num_id") is None:
            list_state["num_id"] = _create_numbering_instance(
                document,
                ordered=False,
            )
        paragraph = document.add_paragraph(style="List Bullet")
        _set_numbering_instance(paragraph, int(list_state["num_id"]))
        paragraph.add_run(text)
        list_state["kind"] = "bullet"
        return

    list_state["kind"] = None
    list_state["num_id"] = None
    if role in {"title", "heading1", "heading2"}:
        level = 2 if role == "heading2" else 1
        document.add_heading(text, level=level)
        return
    paragraph = document.add_paragraph()
    if role == "formula":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.name = "Cambria Math"
        run.italic = True
    else:
        paragraph.add_run(text)


def _remove_initial_empty_paragraph(document: Document) -> None:
    if len(document.paragraphs) != 1 or document.paragraphs[0].text:
        return
    paragraph = document.paragraphs[0]._element
    paragraph.getparent().remove(paragraph)


def _add_block(
    document: Document,
    block: Any,
    *,
    list_state: dict[str, int | None | str] | None = None,
) -> None:
    label = str(_field(block, "block_label") or "text")
    content = str(_field(block, "block_content") or "").strip()
    if not content:
        return
    if label == "table" or "<table" in content.lower():
        if list_state is not None:
            list_state["kind"] = None
            list_state["num_id"] = None
        _add_table(document, content)
    elif label in {"doc_title", "paragraph_title", "title", "figure_title"}:
        if list_state is not None:
            list_state["kind"] = None
            list_state["num_id"] = None
        document.add_heading(_plain_text(content), level=1)
    else:
        _add_content_lines(document, content, label, list_state=list_state)


def _add_ocr_pages(
    document: Document,
    result_list: list[Any],
    page_sizes: list[tuple[float, float]] | None = None,
) -> None:
    for page_index, result in enumerate(result_list):
        if page_index > 0:
            if page_sizes is None:
                document.add_page_break()
            else:
                section = document.add_section(WD_SECTION.NEW_PAGE)
                width_inches, height_inches = page_sizes[
                    min(page_index, len(page_sizes) - 1)
                ]
                _set_section_page(section, width_inches, height_inches, 0.75)
        elif page_sizes:
            width_inches, height_inches = page_sizes[0]
            _set_section_page(
                document.sections[0], width_inches, height_inches, 0.75
            )
        list_state: dict[str, int | None | str] = {"kind": None, "num_id": None}
        document.add_heading(f"第 {page_index + 1} 页", level=2)
        for block in _ordered_blocks(result):
            _add_block(document, block, list_state=list_state)


def _pdf_page_sizes(source_pdf: Path) -> list[tuple[float, float]]:
    reader = PdfReader(str(source_pdf))
    return [
        (
            max(float(page.mediabox.width) / 72.0, 0.01),
            max(float(page.mediabox.height) / 72.0, 0.01),
        )
        for page in reader.pages
    ]


def _add_text_pages_with_sizes(
    document: Document,
    page_texts: Iterable[str],
    page_sizes: list[tuple[float, float]] | None,
) -> int:
    _remove_initial_empty_paragraph(document)
    list_state: dict[str, int | None | str] = {"kind": None, "num_id": None}
    page_count = 0
    for page_index, text in enumerate(page_texts):
        if page_index > 0:
            if page_sizes is None:
                document.add_page_break()
            else:
                section = document.add_section(WD_SECTION.NEW_PAGE)
                width_inches, height_inches = page_sizes[
                    min(page_index, len(page_sizes) - 1)
                ]
                _set_section_page(section, width_inches, height_inches, 0.75)
        elif page_sizes:
            width_inches, height_inches = page_sizes[0]
            _set_section_page(
                document.sections[0], width_inches, height_inches, 0.75
            )
        for role, block_text in _group_text_lines(
            text,
            is_document_start=page_index == 0,
        ):
            _add_editable_text_block(document, role, block_text, list_state)
        page_count += 1
    return page_count


def export_text_pages_to_docx(
    page_texts: Iterable[str],
    output_path: Path,
    title: str,
    *,
    source_pdf: Path | None = None,
    stage_callback: StageCallback | None = None,
) -> None:
    """将 PDF 文本层逐页导出为可编辑 DOCX。"""
    if source_pdf is not None:
        source_pdf = Path(source_pdf)
        if not source_pdf.is_file():
            raise FileNotFoundError(f"文本导出所需的 PDF 不存在：{source_pdf}")
    document = Document()
    _set_document_styles(document)
    document.core_properties.title = title
    text_pages = list(page_texts)
    _notify_stage(stage_callback, "text_export_started", page_count=len(text_pages))
    page_sizes = _pdf_page_sizes(source_pdf) if source_pdf is not None else None
    page_count = _add_text_pages_with_sizes(document, text_pages, page_sizes)
    _notify_stage(stage_callback, "text_export_pages_completed", page_count=page_count)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    _notify_stage(stage_callback, "text_export_completed", page_count=page_count)


def export_source_pages_to_docx(
    output_path: Path,
    title: str,
    *,
    source_pdf: Path,
    image_max_pixels: int = DEFAULT_PAGE_IMAGE_MAX_PIXELS,
    image_jpeg_quality: int = DEFAULT_PAGE_IMAGE_JPEG_QUALITY,
    stage_callback: StageCallback | None = None,
) -> None:
    """按原始 PDF 页面渲染结果生成视觉保真的 DOCX。"""
    if not source_pdf.is_file():
        raise FileNotFoundError(f"页面导出所需的 PDF 不存在：{source_pdf}")
    if image_max_pixels < 1:
        raise ValueError("页面图像像素上限必须大于 0")

    document = Document()
    _set_document_styles(document)
    document.core_properties.title = title
    _notify_stage(stage_callback, "page_images_started")
    image_page_count = _add_source_pages(
        document,
        source_pdf,
        max_pixels=image_max_pixels,
        jpeg_quality=image_jpeg_quality,
        stage_callback=stage_callback,
    )
    _notify_stage(
        stage_callback,
        "page_images_completed",
        page_count=image_page_count,
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    _notify_stage(
        stage_callback,
        "source_pages_completed",
        page_count=image_page_count,
    )


def export_results_to_docx(
    results: Iterable[Any],
    output_path: Path,
    title: str,
    *,
    source_pdf: Path | None = None,
    mode: str = "text",
    image_max_pixels: int = DEFAULT_PAGE_IMAGE_MAX_PIXELS,
    image_jpeg_quality: int = DEFAULT_PAGE_IMAGE_JPEG_QUALITY,
    stage_callback: StageCallback | None = None,
) -> None:
    """将多页 PaddleOCR 版面结果导出为文字版或混合版 DOCX。"""
    normalized_mode = mode.strip().lower()
    if normalized_mode not in EXPORT_MODES:
        raise ValueError(f"不支持的 DOCX 导出模式：{mode}")
    if image_max_pixels < 1:
        raise ValueError("页面图像像素上限必须大于 0")
    if source_pdf is not None:
        source_pdf = Path(source_pdf)
    document = Document()
    _set_document_styles(document)
    document.core_properties.title = title
    result_list = list(results)

    if normalized_mode == "hybrid":
        if source_pdf is None or not source_pdf.is_file():
            raise FileNotFoundError(f"混合导出所需的 PDF 不存在：{source_pdf}")
        _notify_stage(
            stage_callback,
            "page_images_started",
            page_count=len(result_list),
        )
        image_page_count = _add_source_pages(
            document,
            source_pdf,
            max_pixels=image_max_pixels,
            jpeg_quality=image_jpeg_quality,
            stage_callback=stage_callback,
        )
        _notify_stage(
            stage_callback,
            "page_images_completed",
            page_count=image_page_count,
        )
        ocr_section = document.add_section(WD_SECTION.NEW_PAGE)
        _set_section_page(ocr_section, 8.5, 11.0, 0.75)
        _notify_stage(
            stage_callback,
            "ocr_text_started",
            page_count=len(result_list),
        )
        _add_ocr_pages(document, result_list)
        _notify_stage(
            stage_callback,
            "ocr_text_completed",
            page_count=len(result_list),
        )
    else:
        page_sizes = _pdf_page_sizes(source_pdf) if source_pdf is not None else None
        _add_ocr_pages(document, result_list, page_sizes)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
